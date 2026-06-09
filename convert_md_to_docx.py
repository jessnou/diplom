#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

md_path = '/home/jessnou/projects/diplom/word-docs/04b_Обзор_системы.md'
docx_path = '/home/jessnou/projects/diplom/word-docs/04b_Обзор_системы.docx'

doc = Document()

# Page settings
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Heading styles
for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    sizes = {1: 16, 2: 15, 3: 14}
    h_style.font.size = Pt(sizes[level])
    h_style.font.bold = True

with open(md_path, 'r', encoding='utf-8') as f:
    text = f.read()

# Split into blocks separated by empty lines
blocks = re.split(r'\n\n+', text)

in_code = False
code_buffer = []

for block in blocks:
    block = block.strip()
    if not block:
        continue

    lines_in_block = block.split('\n')

    # Check for code block fence
    if lines_in_block[0].startswith('```'):
        if in_code:
            # End code block - add all buffered lines as code paragraphs
            for cl in code_buffer:
                p = doc.add_paragraph()
                run = p.add_run(cl)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            code_buffer = []
            in_code = False
        else:
            in_code = True
        continue

    if in_code:
        code_buffer.extend(lines_in_block)
        continue

    # Heading: starts with # (first line of block)
    first_line = lines_in_block[0]
    heading_match = re.match(r'^(#{1,3})\s+(.+)$', first_line)
    if heading_match and len(lines_in_block) == 1:
        level = len(heading_match.group(1))
        title = heading_match.group(2)
        doc.add_heading(title, level=min(level, 3))
        continue

    # Separator ---
    if first_line == '---':
        continue

    # Table: starts with |
    if first_line.startswith('|') and first_line.endswith('|'):
        # Parse table rows, skip separator
        table_rows = []
        for line in lines_in_block:
            line = line.strip()
            if not line.startswith('|'):
                continue
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue  # separator row
            cells = [c.strip() for c in line.strip('|').split('|')]
            table_rows.append(cells)

        if table_rows:
            cols = max(len(row) for row in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = 'Table Grid'
            for row_idx, row_data in enumerate(table_rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= cols:
                        break
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(cell_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if row_idx == 0:
                        run.bold = True
                        shading = cell._element.get_or_add_tcPr()
                        shd = shading.makeelement(qn('w:shd'), {
                            qn('w:fill'): 'D9D9D9',
                            qn('w:val'): 'clear'
                        })
                        shading.append(shd)
            doc.add_paragraph()
        continue

    # Regular paragraph with potential inline **bold**
    paragraph_text = block.replace('\n', ' ')
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.+?\*\*)', paragraph_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

doc.save(docx_path)
print(f'Saved: {docx_path}')
